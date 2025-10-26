"""
주택임대차 표준계약서 생성 Tool
DOCX 템플릿 기반으로 주택임대차 계약서를 생성
"""

import logging
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class LeaseContractGeneratorTool:
    """
    주택임대차 표준계약서 생성 도구
    DOCX 템플릿을 로드하여 표(table) 셀에 값을 채워 계약서 생성
    """

    def __init__(self, template_path: Optional[str] = None):
        """초기화"""
        self.name = "lease_contract_generator"

        # 기본 템플릿 경로 (플레이스홀더 버전 우선 사용)
        if template_path:
            self.template_path = Path(template_path)
        else:
            backend_dir = Path(__file__).parent.parent.parent.parent
            # 플레이스홀더 버전 템플릿 경로
            placeholder_template = backend_dir / "data" / "storage" / "documents" / "주택임대차 표준계약서_with_placeholders.docx"
            original_template = backend_dir / "data" / "storage" / "documents" / "주택임대차 표준계약서.docx"

            # 플레이스홀더 버전이 있으면 우선 사용, 없으면 원본 사용
            if placeholder_template.exists():
                self.template_path = placeholder_template
            else:
                self.template_path = original_template

        # python-docx 가용성 확인
        self.docx_available = self._check_docx_availability()
        logger.info(f"LeaseContractGeneratorTool initialized (docx: {self.docx_available})")

    def _check_docx_availability(self) -> bool:
        """python-docx 라이브러리 가용성 확인"""
        try:
            import docx
            return True
        except ImportError:
            return False

    async def execute(
        self,
        # 물건 정보
        address_road: Optional[str] = None,
        address_detail: Optional[str] = None,
        land_area: Optional[str] = None,
        building_area: Optional[str] = None,
        rental_area: Optional[str] = None,
        # 금액
        deposit: Optional[str] = None,
        deposit_hangeul: Optional[str] = None,
        contract_payment: Optional[str] = None,
        monthly_rent: Optional[str] = None,
        monthly_rent_day: Optional[str] = None,
        management_fee: Optional[str] = None,
        # 기간
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
        # 당사자
        lessor_name: Optional[str] = None,
        lessor_address: Optional[str] = None,
        lessor_phone: Optional[str] = None,
        lessee_name: Optional[str] = None,
        lessee_address: Optional[str] = None,
        lessee_phone: Optional[str] = None,
        # 특약
        special_terms: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """주택임대차 계약서 생성"""
        try:
            if not self.docx_available:
                return self._return_docx_unavailable_error()

            if not self.template_path.exists():
                return self._return_template_not_found_error()

            # docx 로드
            from docx import Document

            try:
                doc = Document(str(self.template_path))
            except Exception as e:
                logger.error(f"Failed to load docx: {e}")
                return self._return_template_load_error(e)

            # 필드 준비
            fields = self._build_fields(locals())

            # 표 셀 채우기
            self._fill_tables(doc, fields)

            # 출력 파일 저장
            output_path = self._get_output_path()
            doc.save(str(output_path))

            # Markdown 생성
            markdown = self._to_markdown(doc, fields)

            return {
                "status": "success",
                "content": markdown,
                "title": "주택임대차 표준계약서",
                "docx_path": str(output_path),
                "fields": fields,
                "sections": self._extract_sections(markdown),
                "metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "output_docx": str(output_path)
                },
                "timestamp": datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"Contract generation failed: {e}", exc_info=True)
            return {
                "status": "error",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }

    def _build_fields(self, params: Dict) -> Dict:
        """필드 딕셔너리 구성"""
        fields = {k: v for k, v in params.items() if k not in ['self', 'kwargs'] and v}

        # 기본값
        defaults = {
            "address_road": "[도로명주소]",
            "deposit": "[보증금]",
            "start_date": "[시작일]",
            "end_date": "[종료일]",
            "lessor_name": "[임대인명]",
            "lessee_name": "[임차인명]"
        }

        for k, v in defaults.items():
            if k not in fields or not fields[k]:
                fields[k] = v

        return fields

    def _fill_tables(self, doc, fields: Dict):
        """표 셀 채우기 (플레이스홀더 기반)"""
        try:
            if len(doc.tables) == 0:
                return

            # 플레이스홀더 매핑 정의
            placeholder_map = {
                "{{address_road}}": fields.get("address_road", ""),
                "{{address_detail}}": fields.get("address_detail", ""),
                "{{rental_area}}": fields.get("rental_area", ""),
                "{{deposit}}": fields.get("deposit", ""),
                "{{deposit_hangeul}}": fields.get("deposit_hangeul", ""),
                "{{contract_payment}}": fields.get("contract_payment", ""),
                "{{monthly_rent}}": fields.get("monthly_rent", ""),
                "{{monthly_rent_day}}": fields.get("monthly_rent_day", ""),
                "{{management_fee}}": fields.get("management_fee", ""),
                "{{lessor_name}}": fields.get("lessor_name", ""),
                "{{lessor_address}}": fields.get("lessor_address", ""),
                "{{lessor_phone}}": fields.get("lessor_phone", ""),
                "{{lessee_name}}": fields.get("lessee_name", ""),
                "{{lessee_address}}": fields.get("lessee_address", ""),
                "{{lessee_phone}}": fields.get("lessee_phone", ""),
            }

            # 모든 테이블 순회하며 플레이스홀더 치환
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        new_text = original_text

                        # 플레이스홀더 치환
                        for placeholder, value in placeholder_map.items():
                            if placeholder in new_text and value:
                                new_text = new_text.replace(placeholder, str(value))

                        # 텍스트 변경된 경우만 업데이트
                        if new_text != original_text:
                            cell.text = new_text

            logger.info("Placeholders filled successfully")

        except Exception as e:
            logger.error(f"Failed to fill tables: {e}")


    def _to_markdown(self, doc, fields: Dict) -> str:
        """Markdown 변환"""
        lines = ["# 주택임대차 표준계약서\n"]
        lines.append(f"**작성일**: {datetime.now().strftime('%Y년 %m월 %d일')}\n")

        lines.append("## 물건 정보")
        lines.append(f"- **소재지**: {fields.get('address_road', '[주소]')}")
        lines.append(f"- **상세주소**: {fields.get('address_detail', '[동/층/호]')}")
        lines.append(f"- **임차면적**: {fields.get('rental_area', '[면적]')} ㎡\n")

        lines.append("## 계약 금액")
        lines.append(f"- **보증금**: {fields.get('deposit', '[보증금]')} 원")
        if fields.get('monthly_rent'):
            lines.append(f"- **월세**: {fields['monthly_rent']} 원")
        if fields.get('management_fee'):
            lines.append(f"- **관리비**: {fields['management_fee']} 원\n")

        lines.append("## 임대차 기간")
        lines.append(f"- **시작일**: {fields.get('start_date', '[시작일]')}")
        lines.append(f"- **종료일**: {fields.get('end_date', '[종료일]')}\n")

        lines.append("## 당사자")
        lines.append(f"- **임대인**: {fields.get('lessor_name', '[임대인명]')}")
        if fields.get('lessor_phone'):
            lines.append(f"  - 연락처: {fields['lessor_phone']}")
        lines.append(f"- **임차인**: {fields.get('lessee_name', '[임차인명]')}")
        if fields.get('lessee_phone'):
            lines.append(f"  - 연락처: {fields['lessee_phone']}\n")

        if fields.get('special_terms'):
            lines.append("## 특약사항")
            lines.append(f"{fields['special_terms']}\n")

        return "\n".join(lines)

    def _extract_sections(self, content: str) -> list:
        """섹션 추출"""
        sections = []
        current_section = None
        current_content = []

        for line in content.split("\n"):
            if line.strip().startswith("#"):
                if current_section:
                    sections.append({
                        "title": current_section,
                        "content": "\n".join(current_content).strip()
                    })
                current_section = line.strip().replace("#", "").strip()
                current_content = []
            elif current_section:
                current_content.append(line)

        if current_section:
            sections.append({
                "title": current_section,
                "content": "\n".join(current_content).strip()
            })

        return sections

    def _get_output_path(self) -> Path:
        """출력 파일 경로"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = self.template_path.parent / "generated"
        output_dir.mkdir(exist_ok=True)
        return output_dir / f"주택임대차계약서_{timestamp}.docx"

    def _return_docx_unavailable_error(self) -> Dict:
        """python-docx 없음 에러"""
        return {
            "status": "error",
            "error_type": "docx_unavailable",
            "message": (
                "python-docx 라이브러리가 설치되어 있지 않습니다.\n"
                "설치 명령: pip install python-docx"
            ),
            "timestamp": datetime.now().isoformat()
        }

    def _return_template_not_found_error(self) -> Dict:
        """템플릿 없음 에러"""
        return {
            "status": "error",
            "error_type": "template_not_found",
            "message": (
                f"주택임대차 표준계약서 템플릿을 찾을 수 없습니다.\n"
                f"템플릿 경로: {self.template_path}\n\n"
                f"템플릿 파일을 해당 위치에 배치하거나,\n"
                f".txt 또는 .md 형식으로 변환해주세요."
            ),
            "template_path": str(self.template_path),
            "timestamp": datetime.now().isoformat()
        }

    def _return_template_load_error(self, error: Exception) -> Dict:
        """템플릿 로드 실패 에러"""
        return {
            "status": "error",
            "error_type": "template_load_failed",
            "message": (
                f"주택임대차 표준계약서 템플릿을 로드할 수 없습니다.\n"
                f"오류: {str(error)}\n\n"
                f"템플릿 파일 ({self.template_path})을 확인하거나,\n"
                f".txt 또는 .md 형식으로 변환해주세요."
            ),
            "template_path": str(self.template_path),
            "error": str(error),
            "timestamp": datetime.now().isoformat()
        }

    def get_required_fields(self) -> list:
        """필수 필드"""
        return ["address_road", "deposit", "start_date", "end_date", "lessor_name", "lessee_name"]

    def get_optional_fields(self) -> list:
        """선택 필드"""
        return [
            "address_detail", "land_area", "building_area", "rental_area",
            "deposit_hangeul", "contract_payment", "monthly_rent", "monthly_rent_day",
            "management_fee", "lessor_address", "lessor_phone",
            "lessee_address", "lessee_phone", "special_terms"
        ]


# 테스트
if __name__ == "__main__":
    import asyncio

    async def test():
        tool = LeaseContractGeneratorTool()

        params = {
            "address_road": "서울특별시 강남구 테헤란로 123",
            "address_detail": "456호 (101동 10층)",
            "rental_area": "85",
            "deposit": "500,000,000",
            "deposit_hangeul": "오억",
            "contract_payment": "50,000,000",
            "monthly_rent": "2,000,000",
            "monthly_rent_day": "1",
            "management_fee": "150,000",
            "start_date": "2024년 1월 1일",
            "end_date": "2026년 1월 1일",
            "lessor_name": "홍길동",
            "lessor_address": "서울시 서초구",
            "lessor_phone": "010-1234-5678",
            "lessee_name": "김철수",
            "lessee_address": "서울시 강남구",
            "lessee_phone": "010-9876-5432",
            "special_terms": "반려동물 사육 가능"
        }

        result = await tool.execute(**params)

        print("=== Result ===")
        print(f"Status: {result['status']}")
        if result['status'] == 'success':
            print(f"DOCX: {result.get('docx_path')}")
            print(f"Sections: {len(result.get('sections', []))}")
            print(f"\nPreview:\n{result['content'][:500]}...")
        else:
            print(f"Error: {result.get('error') or result.get('message')}")

    asyncio.run(test())
